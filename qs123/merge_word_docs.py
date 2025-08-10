import os
from docx import Document


def copy_paragraph_style(source_paragraph, target_paragraph):
    """
    将源段落的样式复制到目标段落。
    注意：这不会复制段落的样式名称，而是复制具体的格式设置。
    """
    target_paragraph.style = source_paragraph.style
    target_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
    target_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
    target_paragraph.paragraph_format.keep_together = source_paragraph.paragraph_format.keep_together
    target_paragraph.paragraph_format.keep_with_next = source_paragraph.paragraph_format.keep_with_next
    target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
    target_paragraph.paragraph_format.page_break_before = source_paragraph.paragraph_format.page_break_before
    target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.widow_control = source_paragraph.paragraph_format.widow_control


def copy_run_style(source_run, target_run):
    """
    将源run的字符样式复制到目标run。
    """
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    # 可以根据需要添加更多格式属性的复制


def merge_word_documents(input_folder, output_file):
    """
    合并指定文件夹中的所有Word文档 (.docx) 到一个新文档中。

    :param input_folder: 包含要合并的Word文档的文件夹路径。
    :param output_file: 合并后输出的Word文档的文件名（含路径）。
    """
    # 创建一个新的空文档用于存放合并后的内容
    merged_document = Document()

    # 遍历输入文件夹中的所有.docx文件
    for filename in sorted(os.listdir(input_folder)):  # 使用sorted确保按文件名顺序处理
        if filename.endswith('.docx'):
            file_path = os.path.join(input_folder, filename)
            print(f"正在处理文件: {filename}")

            # 打开要合并的文档
            doc_to_merge = Document(file_path)

            # 为每个新文档添加一个分节符（下一页），以保持原有页面布局
            if merged_document.paragraphs:  # 如果不是第一个文档，则添加分节符
                merged_document.add_section()

            # 遍历要合并文档中的所有段落，并添加到新文档中
            for paragraph in doc_to_merge.paragraphs:
                # 添加一个新的段落到合并后的文档
                new_paragraph = merged_document.add_paragraph()
                # 复制段落的文本和基本格式
                new_paragraph.text = paragraph.text
                copy_paragraph_style(paragraph, new_paragraph)

                # 清除默认的run，因为我们将在下面重新创建它们以保留格式
                for run in new_paragraph.runs:
                    run._element.getparent().remove(run._element)

                # 重新创建run并复制详细的字符格式
                for run_orig in paragraph.runs:
                    new_run = new_paragraph.add_run(run_orig.text)
                    copy_run_style(run_orig, new_run)

            # 复制表格
            for table in doc_to_merge.tables:
                # 在合并文档中添加一个新表格，行列数与原表格相同
                new_table = merged_document.add_table(rows=len(table.rows), cols=len(table.columns))
                # 复制表格内容和基本格式
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
                        # 注意：更复杂的表格格式（如边框、底纹）复制较为复杂，这里未实现

    # 保存合并后的文档
    merged_document.save(output_file)
    print(f"所有文档已成功合并到 {output_file}")


if __name__ == "__main__":
    # 指定输入文件夹和输出文件名
    input_folder_path = r"C:\Users\hk\word_files_to_merge"  # 替换为你的Word文件所在文件夹
    output_file_path = r"C:\Users\hk\merged_document_with_formatting.docx"  # 替换为你想保存的合并后文件的路径

    merge_word_documents(input_folder_path, output_file_path)