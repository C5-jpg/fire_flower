from docx import Document

# 创建第一个示例文档
doc1 = Document()
doc1.add_heading('文档一', level=1)
doc1.add_paragraph('这是第一个示例文档的内容。')
doc1.add_paragraph('它包含一些文本。', style='Intense Quote')
doc1.save(r'C:\Users\hk\word_files_to_merge\document1.docx')

# 创建第二个示例文档
doc2 = Document()
doc2.add_heading('文档二', level=1)
doc2.add_paragraph('这是第二个示例文档的内容。')
doc2.add_paragraph('它也包含一些文本，但格式略有不同。', style='List Bullet')
doc2.save(r'C:\Users\hk\word_files_to_merge\document2.docx')

# 创建第三个示例文档
doc3 = Document()
doc3.add_heading('文档三', level=1)
doc3.add_paragraph('这是第三个示例文档的内容。')
doc3.add_paragraph('这个文档用来测试合成功能。')
table = doc3.add_table(rows=2, cols=2)
table.cell(0, 0).text = '表头1'
table.cell(0, 1).text = '表头2'
table.cell(1, 0).text = '数据1'
table.cell(1, 1).text = '数据2'
doc3.save(r'C:\Users\hk\word_files_to_merge\document3.docx')